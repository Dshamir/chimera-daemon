"""Document extractors (PDF, DOCX, MD, TXT, HTML)."""

from pathlib import Path

from chimera.extractors.base import BaseExtractor, ExtractionResult
from chimera.utils.logging import get_logger

logger = get_logger(__name__)


class PDFExtractor(BaseExtractor):
    """Extract content from PDF files."""
    
    name = "pdf"
    extensions = ["pdf"]
    mime_types = ["application/pdf"]
    
    async def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from PDF."""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            content = "\n\n".join(text_parts)
            
            return ExtractionResult(
                file_path=file_path,
                content=content,
                metadata={
                    "page_count": len(reader.pages),
                    "pdf_info": dict(reader.metadata) if reader.metadata else {},
                },
                word_count=self.count_words(content),
                page_count=len(reader.pages),
            )
        except Exception as e:
            logger.error(f"PDF extraction failed: {file_path}: {e}")
            return ExtractionResult(
                file_path=file_path,
                content="",
                success=False,
                error=str(e),
            )


class DOCXExtractor(BaseExtractor):
    """Extract content from DOCX files."""
    
    name = "docx"
    extensions = ["docx"]
    mime_types = ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
    
    async def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from DOCX."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    tables_text.append(row_text)
            
            if tables_text:
                content += "\n\n[Tables]\n" + "\n".join(tables_text)
            
            return ExtractionResult(
                file_path=file_path,
                content=content,
                metadata={
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables),
                },
                word_count=self.count_words(content),
            )
        except Exception as e:
            logger.error(f"DOCX extraction failed: {file_path}: {e}")
            return ExtractionResult(
                file_path=file_path,
                content="",
                success=False,
                error=str(e),
            )


class MarkdownExtractor(BaseExtractor):
    """Extract content from Markdown files."""
    
    name = "markdown"
    extensions = ["md", "markdown"]
    mime_types = ["text/markdown"]
    
    async def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from Markdown."""
        try:
            content = file_path.read_text(encoding="utf-8")
            
            # Extract headers for metadata
            headers = []
            for line in content.split("\n"):
                if line.startswith("#"):
                    level = len(line) - len(line.lstrip("#"))
                    text = line.lstrip("# ").strip()
                    headers.append({"level": level, "text": text})
            
            return ExtractionResult(
                file_path=file_path,
                content=content,
                metadata={
                    "format": "markdown",
                    "headers": headers,
                    "header_count": len(headers),
                },
                word_count=self.count_words(content),
            )
        except Exception as e:
            logger.error(f"Markdown extraction failed: {file_path}: {e}")
            return ExtractionResult(
                file_path=file_path,
                content="",
                success=False,
                error=str(e),
            )


class TextExtractor(BaseExtractor):
    """Extract content from plain text files."""
    
    name = "text"
    extensions = ["txt", "text", "log"]
    mime_types = ["text/plain"]
    
    async def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from plain text file."""
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                content = file_path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                content = file_path.read_text(encoding="latin-1")
            
            return ExtractionResult(
                file_path=file_path,
                content=content,
                metadata={
                    "format": "text",
                    "line_count": len(content.splitlines()),
                },
                word_count=self.count_words(content),
            )
        except Exception as e:
            logger.error(f"Text extraction failed: {file_path}: {e}")
            return ExtractionResult(
                file_path=file_path,
                content="",
                success=False,
                error=str(e),
            )


class HTMLExtractor(BaseExtractor):
    """Extract content from HTML files."""
    
    name = "html"
    extensions = ["html", "htm"]
    mime_types = ["text/html"]
    
    async def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            
            html_content = file_path.read_text(encoding="utf-8")
            soup = BeautifulSoup(html_content, "html.parser")
            
            # Remove script and style elements
            for element in soup(["script", "style", "nav", "footer"]):
                element.decompose()
            
            # Get text
            content = soup.get_text(separator="\n", strip=True)
            
            # Extract title
            title = soup.title.string if soup.title else None
            
            # Extract links
            links = [a.get("href") for a in soup.find_all("a", href=True)]
            
            return ExtractionResult(
                file_path=file_path,
                content=content,
                metadata={
                    "format": "html",
                    "title": title,
                    "link_count": len(links),
                },
                word_count=self.count_words(content),
            )
        except Exception as e:
            logger.error(f"HTML extraction failed: {file_path}: {e}")
            return ExtractionResult(
                file_path=file_path,
                content="",
                success=False,
                error=str(e),
            )
